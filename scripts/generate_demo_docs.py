from __future__ import annotations

import json
from pathlib import Path
import sys

WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))

from docx import Document
from docx.shared import Inches
from win32com.client import Dispatch


SCREENSHOT_ORDER = [
    ("project_center", "项目中心：确认演示项目、零件基础信息和项目目录。"),
    ("parts", "零件与材质：展示零件尺寸、质量要求以及材质库热物性数据。"),
    ("scheme", "工艺方案：查看基础优化方案、分型方式和浇注系统设置。"),
    ("parameters", "参数计算：展示关键参数、公式依据和计算结果。"),
    ("solidworks", "SolidWorks 协同：关联本地三维模型并识别桥接状态。"),
    ("simulation", "ProCAST 仿真：查看仿真任务登记、输入输出目录和任务状态。"),
    ("results", "结果对比：展示缩孔/温度类结果索引和对比组。"),
    ("documents", "工艺卡与质检：显示已生成的工艺卡与清单。"),
    ("ai", "AI 助手：展示自动生成的建议卡片和证据。"),
    ("review", "建议审核：保留待审核建议，体现人工确认环节。"),
    ("export", "成果导出：展示自动识别到的最新成果包目录和 ZIP。"),
    ("settings", "系统设置：展示 LLM、桥接路径、ProCAST 路径和默认项目目录配置。"),
]


def main() -> None:
    workspace = Path(__file__).resolve().parents[1]
    artifacts_dir = workspace / "artifacts"
    walkthrough_dir = artifacts_dir / "walkthrough"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    manual_md = artifacts_dir / "砂型铸造本地智能工作站_操作说明书.md"
    manual_docx = artifacts_dir / "砂型铸造本地智能工作站_操作说明书.docx"
    ppt_outline = artifacts_dir / "砂型铸造本地智能工作站_PPT提纲.md"
    pptx_path = artifacts_dir / "砂型铸造本地智能工作站_创新课题材料.pptx"

    manual_md.write_text(_build_manual_markdown(summary), encoding="utf-8")
    ppt_outline.write_text(_build_ppt_outline(summary), encoding="utf-8")
    _build_manual_docx(summary, walkthrough_dir, manual_docx)
    _build_ppt(summary, walkthrough_dir, pptx_path)

    manifest_path = artifacts_dir / "demo_materials_manifest.md"
    manifest_path.write_text(
        "\n".join(
            [
                "# 演示材料清单",
                "",
                f"- 操作说明书 Markdown：{manual_md}",
                f"- 操作说明书 Word：{manual_docx}",
                f"- PPT 提纲：{ppt_outline}",
                f"- 创新课题 PPT：{pptx_path}",
                f"- 操作视频：{artifacts_dir / 'demo_walkthrough.mp4'}",
                "",
                "## 关键截图",
                *[f"- {walkthrough_dir / f'{name}.png'}" for name, _ in SCREENSHOT_ORDER],
                "",
            ]
        ),
        encoding="utf-8",
    )


def _build_manual_markdown(summary: dict) -> str:
    lines = [
        "# 砂型铸造本地智能工作站操作说明书",
        "",
        "## 1. 演示目标",
        "完成一次从项目管理、工艺方案、参数计算、CAD 协同、仿真结果管理到工艺卡、AI 建议与成果包导出的完整闭环演示。",
        "",
        "## 2. 演示环境",
        f"- 项目编号：{summary['project_code']}",
        f"- 项目名称：{summary['project_name']}",
        f"- 项目目录：{summary['project_root']}",
        "- 本地集成：SolidWorks、ProCAST 2022.0",
        "- 桌面程序：CastingWorkstation 便携版",
        "",
        "## 3. 演示结果",
        f"- 工艺卡：{summary['document_card_path']}",
        f"- 质检/缺陷预防清单：{summary['checklist_path']}",
        f"- AI 建议数：{summary['suggestion_count']}",
        f"- 待审核建议数：{summary['pending_review_count']}",
        f"- 成果包 ZIP：{summary['export_zip_path']}",
        "",
        "## 4. 标准操作步骤",
    ]
    for index, (_, caption) in enumerate(SCREENSHOT_ORDER, start=1):
        lines.extend([f"{index}. {caption}"])
    lines.extend(
        [
            "",
            "## 5. 创新点",
            "- 本地桌面工作站为主，适配教学和竞赛场景。",
            "- 统一管理项目、工艺参数、CAD 文件、仿真结果和工艺卡。",
            "- 在 AI 建议中保留证据链和人工审核状态，避免黑箱输出。",
            "- 预留远程同步、教师查看和网页只读展示接口。",
            "",
            "## 6. 后续扩展",
            "- 接入真实 LLM 推理服务，替换当前规则型建议生成。",
            "- 增加 ProCAST 任务模板自动生成和结果缩略图预览。",
            "- 增加工艺图模板、答辩报告模板和教学评价看板。",
            "",
        ]
    )
    return "\n".join(lines)


def _build_ppt_outline(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站 PPT 提纲",
            "",
            "1. 课题背景与目标",
            "2. 系统总体架构与技术路线",
            "3. 演示案例与运行环境",
            "4. 项目中心与工艺方案管理",
            "5. 参数计算、CAD/CAE 协同与结果管理",
            "6. 工艺卡、AI 建议与审核闭环",
            "7. 成果包导出与当前演示结果",
            "8. 创新点、价值与后续计划",
            "",
            f"- 演示项目：{summary['project_code']} / {summary['project_name']}",
            f"- 成果包：{summary['export_zip_path']}",
        ]
    )


def _build_manual_docx(summary: dict, walkthrough_dir: Path, output_path: Path) -> None:
    document = Document()
    document.add_heading("砂型铸造本地智能工作站操作说明书", level=0)
    document.add_paragraph("用途：用于创新课题演示、竞赛汇报和软件交付说明。")

    document.add_heading("一、演示目标", level=1)
    document.add_paragraph(
        "本次演示围绕汽油机端盖案例，验证系统在项目管理、工艺方案、参数计算、SolidWorks 协同、"
        "ProCAST 仿真结果管理、工艺卡生成、AI 建议和成果包导出上的完整闭环能力。"
    )

    document.add_heading("二、演示环境", level=1)
    for line in (
        f"项目编号：{summary['project_code']}",
        f"项目名称：{summary['project_name']}",
        f"项目目录：{summary['project_root']}",
        "本机已安装 SolidWorks 与 ProCAST 2022.0。",
        "演示程序采用本地便携版 CastingWorkstation。",
    ):
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("三、关键输出", level=1)
    for line in (
        f"工艺卡：{summary['document_card_path']}",
        f"质检/缺陷预防清单：{summary['checklist_path']}",
        f"AI 建议数：{summary['suggestion_count']}",
        f"待审核建议数：{summary['pending_review_count']}",
        f"成果包 ZIP：{summary['export_zip_path']}",
    ):
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("四、完整操作过程", level=1)
    for index, (name, caption) in enumerate(SCREENSHOT_ORDER, start=1):
        document.add_heading(f"步骤 {index}", level=2)
        document.add_paragraph(caption)
        image_path = walkthrough_dir / f"{name}.png"
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.5))

    document.add_heading("五、创新点总结", level=1)
    for line in (
        "本地桌面工作站优先，便于教学、竞赛与实验室场景部署。",
        "统一数据模型贯通项目、参数、CAD、仿真、工艺卡与 AI 建议。",
        "AI 输出保留证据链与人工审核记录，符合工程审慎原则。",
        "架构天然支持后续扩展教师查看、远程同步与网页只读展示。",
    ):
        document.add_paragraph(line, style="List Bullet")

    document.save(output_path)


def _build_ppt(summary: dict, walkthrough_dir: Path, output_path: Path) -> None:
    app = Dispatch("PowerPoint.Application")
    app.Visible = 1
    presentation = app.Presentations.Add()

    def add_title_slide(title: str, subtitle: str) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 1)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = subtitle

    def add_bullet_slide(title: str, bullets: list[str]) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 2)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        text_range = slide.Shapes.Placeholders(2).TextFrame.TextRange
        text_range.Text = "\r".join(bullets)

    def add_picture_slide(title: str, image_name: str, bullets: list[str]) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 2)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        text_range = slide.Shapes.Placeholders(2).TextFrame.TextRange
        text_range.Text = "\r".join(bullets)
        image_path = walkthrough_dir / f"{image_name}.png"
        if image_path.exists():
            slide.Shapes.AddPicture(str(image_path), False, True, 430, 120, 460, 300)

    add_title_slide(
        "砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站",
        f"创新课题演示材料\n案例：{summary['project_name']}",
    )
    add_bullet_slide(
        "课题背景与目标",
        [
            "面向砂型铸造工艺设计教学、竞赛与实验室应用。",
            "以本地桌面工作站为主，整合参数计算、CAD/CAE 协同、工艺卡输出与 AI 建议。",
            "避免做成 PLC 实时闭环控制系统，强调人工审核和工程证据链。",
        ],
    )
    add_bullet_slide(
        "系统总体架构",
        [
            "PySide6 桌面前端 + Python 本地业务核心。",
            "SQLite 统一管理项目、参数、仿真、文档与建议记录。",
            "SolidWorks 通过本地桥接协同，ProCAST 通过任务/结果目录归档协同。",
        ],
    )
    add_picture_slide(
        "项目中心与工艺方案",
        "project_center",
        [
            f"演示项目：{summary['project_code']}",
            "维护零件、材质、项目目录和工艺方案版本。",
            "为参数计算、文档生成和仿真结果绑定统一上下文。",
        ],
    )
    add_picture_slide(
        "零件与材质页面",
        "parts",
        [
            "展示零件尺寸、壁厚、质量等级、热处理和表面要求。",
            "内置小型材质库，支持查看密度、液相线、浇注温度和收缩率。",
            "为后续工艺参数计算和教师讲解提供统一材料基础数据。",
        ],
    )
    add_picture_slide(
        "参数计算与 CAD/CAE 协同",
        "parameters",
        [
            "自动计算出品率、阻流截面积、浇注时间等关键参数。",
            "本地识别 SolidWorks 模型文件和桥接状态。",
            "仿真任务与结果目录与项目方案绑定。",
        ],
    )
    add_picture_slide(
        "仿真结果与工艺文件",
        "results",
        [
            "统一查看仿真任务、结果索引和对比组。",
            f"工艺卡输出：{Path(summary['document_card_path']).name}",
            "可进一步形成竞赛答辩或教学归档资料。",
        ],
    )
    add_picture_slide(
        "AI 建议与人工审核",
        "ai",
        [
            f"自动生成 {summary['suggestion_count']} 条建议卡片。",
            "建议保留目标参数、风险提示、置信度和证据来源。",
            f"当前仍有 {summary['pending_review_count']} 条建议待人工审核。",
        ],
    )
    add_picture_slide(
        "成果包导出",
        "export",
        [
            "自动汇总工艺卡、质检清单、CAD 文件和仿真结果。",
            f"ZIP 输出：{Path(summary['export_zip_path']).name}",
            "便于竞赛提交、答辩展示和课题归档。",
        ],
    )
    add_picture_slide(
        "系统设置与集成配置",
        "settings",
        [
            "集中管理 LLM 模式、模型名称、SolidWorksBridge 和 ProCAST 路径。",
            "将本地环境配置显式化，便于后续复制到其他计算机运行。",
            "为远程同步、教师查看和网页只读展示预留扩展入口。",
        ],
    )
    add_bullet_slide(
        "创新点与后续计划",
        [
            "本地优先、桌面优先，适合学校实验室和课程项目部署。",
            "统一数据模型贯通工艺设计到成果输出的全过程。",
            "后续继续接入真实 LLM、教师查看端和网页只读展示。",
        ],
    )

    presentation.SaveAs(str(output_path))
    presentation.Close()
    app.Quit()


if __name__ == "__main__":
    main()
