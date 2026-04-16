from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))


def main() -> None:
    artifacts_dir = WORKSPACE / "artifacts"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    body_md = artifacts_dir / "砂型铸造本地智能工作站_申请书正文素材.md"
    body_docx = artifacts_dir / "砂型铸造本地智能工作站_申请书正文素材.docx"
    pitch_md = artifacts_dir / "砂型铸造本地智能工作站_3分钟极简汇报稿.md"
    pitch_docx = artifacts_dir / "砂型铸造本地智能工作站_3分钟极简汇报稿.docx"
    manifest_path = artifacts_dir / "创新训练项目材料清单.md"

    body_text = build_application_body(summary)
    pitch_text = build_short_pitch(summary)

    body_md.write_text(body_text, encoding="utf-8")
    pitch_md.write_text(pitch_text, encoding="utf-8")
    build_docx("砂型铸造本地智能工作站申请书正文素材", body_text, body_docx)
    build_docx("砂型铸造本地智能工作站 3 分钟极简汇报稿", pitch_text, pitch_docx)
    update_manifest(manifest_path, body_md, body_docx, pitch_md, pitch_docx)


def build_application_body(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站申请书正文素材",
            "",
            "## 一、项目简介",
            "本项目拟开发一套“砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站”，以本地桌面系统为核心，围绕铸造工艺设计、工艺卡生成、仿真结果管理和 AI 辅助建议构建统一工作平台。",
            "系统重点面向课程教学、创新训练、学科竞赛和实验室项目管理场景，优先兼容本地安装的 SolidWorks 和 ProCAST，不以纯网页为主，也不设计为工业 PLC 实时闭环控制系统。",
            "",
            "## 二、研究背景与意义",
            "当前砂型铸造教学与竞赛训练中，常见问题是工艺图、工艺卡、三维模型、仿真结果和优化建议分散在不同软件与文件夹中，造成数据难以统一管理、成果难以复用、演示材料难以组织。",
            "随着数字化设计和智能辅助技术的发展，将工艺参数计算、CAD/CAE 协同、文档生成和 AI 建议整合到同一平台，具有明显的教学价值和工程实践价值。",
            "本项目有助于提高学生对铸造工艺设计流程的整体理解，增强参数分析、仿真解读和方案论证能力，同时为实验室构建可复用的本地化数字工具。",
            "",
            "## 三、研究目标",
            "1. 建设本地优先的桌面智能工作站，形成稳定可运行的软件原型。",
            "2. 打通项目管理、零件与材质、工艺方案、参数计算、工艺卡生成和成果导出全流程。",
            "3. 与 SolidWorks、ProCAST 建立协同关系，实现 CAD/CAE 文件与任务归档管理。",
            "4. 引入 AI 建议卡片、证据链和人工审核机制，提升建议的可解释性与可用性。",
            "5. 为后续教师查看、远程同步和网页只读展示预留扩展接口。",
            "",
            "## 四、主要研究内容",
            "1. 设计统一数据模型，覆盖项目、零件、材质、工艺方案、参数、CAD 文件、仿真任务、文档和 AI 建议。",
            "2. 构建本地桌面前端，实现项目中心、零件与材质、工艺方案、参数计算、仿真中心、结果对比、工艺卡、AI 助手、审核与导出页面。",
            "3. 建立参数计算模块，实现浇注时间、阻流截面积、收缩补偿等关键参数计算。",
            "4. 完成与 SolidWorks 的模型关联与导出协同，以及与 ProCAST 的仿真任务与结果目录协同。",
            "5. 实现工艺卡、质检/缺陷预防清单、成果包等自动生成能力。",
            "6. 设计 AI 建议卡片机制，使建议与证据来源、审核结论绑定。",
            "",
            "## 五、技术路线",
            "本项目采用 PySide6 构建桌面前端，Python 作为本地业务核心，SQLite 作为本地数据库。",
            "系统通过统一数据模型串联项目数据、工艺参数、CAD/CAE 资产和文档资产；SolidWorks 负责三维建模和导出，ProCAST 负责仿真，本系统负责组织、管理和输出。",
            "AI 模块采取“建议卡片 + 证据链 + 人工审核”机制，避免直接输出黑箱结论。",
            "",
            "## 六、创新点",
            "1. 本地优先的一体化工作站设计，适合实验室、课程和竞赛现场使用。",
            "2. 将工艺设计、参数计算、仿真结果、工艺卡与 AI 建议统一在同一项目上下文中。",
            "3. AI 建议采用证据化输出与人工审核机制，强调工程可解释性和教学可验证性。",
            "4. 软件运行过程可直接沉淀为成果包和演示材料，便于项目申报与答辩展示。",
            "",
            "## 七、阶段性成果基础",
            f"目前系统已完成一套演示案例，项目编号为 {summary['project_code']}，项目名称为 {summary['project_name']}。",
            "现阶段已跑通以下闭环：",
            f"- 工艺卡生成：{summary['document_card_path']}",
            f"- 质检/缺陷预防清单：{summary['checklist_path']}",
            f"- AI 建议数：{summary['suggestion_count']}",
            f"- 待审核建议数：{summary['pending_review_count']}",
            f"- 成果包 ZIP：{summary['export_zip_path']}",
            "",
            "## 八、预期成果",
            "1. 一套可安装、可运行的本地桌面工作站软件原型。",
            "2. 一套典型零件案例演示数据，包括三维模型、工艺参数、仿真任务和工艺卡。",
            "3. 一套面向答辩和展示的图文说明书、PPT、讲稿和问答提纲。",
            "4. 项目总结报告、阶段性研究资料及相关软件著作权/竞赛成果基础。",
            "",
            "## 九、进度安排",
            "1. 前期阶段：完成需求梳理、总体架构设计、数据库与页面框架搭建。",
            "2. 中期阶段：完成参数计算、CAD/CAE 协同、文档生成和 AI 建议模块开发。",
            "3. 后期阶段：完成联调、演示案例制作、材料整理和答辩准备。",
            "",
            "## 十、应用前景",
            "该系统可直接服务于铸造工艺课程教学、创新训练项目、学科竞赛训练和实验室案例管理。",
            "后续在扩展教师查看端、远程同步与网页只读展示后，还可进一步支撑校内教学资源共享和多角色协作。",
            "",
        ]
    )


def build_short_pitch(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站 3 分钟极简汇报稿",
            "",
            "各位老师好，我们的项目是“砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站”。",
            "这个项目要解决的问题很明确：在铸造教学和竞赛中，工艺图、工艺卡、三维模型和仿真结果通常分散在不同软件里，流程难以打通，成果也不方便展示。",
            "",
            "为了解决这个问题，我们做了一套本地优先的桌面软件，而不是纯网页系统。它能够把项目管理、零件与材质、工艺方案、参数计算、SolidWorks 协同、ProCAST 仿真、工艺卡生成、AI 建议和成果包导出放到同一个平台里。",
            "",
            "在技术实现上，我们采用 PySide6 做桌面界面，Python 做本地业务核心，SQLite 做本地数据库。SolidWorks 和 ProCAST 继续承担专业建模与仿真，本系统负责把这些数据统一组织起来。",
            "",
            f"目前我们已经完成了一套案例演示，项目编号是 {summary['project_code']}。系统已经能够自动生成工艺卡、质检清单、"
            f"{summary['suggestion_count']} 条 AI 建议以及完整成果包。当前还保留 {summary['pending_review_count']} 条待审核建议，用来体现人工确认环节。",
            "",
            "本项目的创新点主要有三点：第一，本地优先，适合实验室和竞赛现场；第二，一体化，把工艺、仿真、文档和 AI 连成闭环；第三，AI 可解释，通过建议卡片、证据链和人工审核提升可信度。",
            "",
            "下一步我们会继续扩充材料库和规则库，接入真实大语言模型，并增加教师查看端和网页只读展示能力。谢谢各位老师。",
            "",
        ]
    )


def build_docx(title: str, text: str, output_path: Path) -> None:
    document = Document()
    lines = text.splitlines()
    if lines:
        document.add_heading(lines[0].lstrip("# ").strip(), level=0)
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(output_path)


def update_manifest(
    manifest_path: Path,
    body_md: Path,
    body_docx: Path,
    pitch_md: Path,
    pitch_docx: Path,
) -> None:
    content = manifest_path.read_text(encoding="utf-8")
    if "## 申请书与极简汇报材料" in content:
        return
    addition = "\n".join(
        [
            "",
            "## 申请书与极简汇报材料",
            f"- 申请书正文素材 Markdown：{body_md}",
            f"- 申请书正文素材 Word：{body_docx}",
            f"- 3 分钟极简汇报稿 Markdown：{pitch_md}",
            f"- 3 分钟极简汇报稿 Word：{pitch_docx}",
            "",
        ]
    )
    manifest_path.write_text(content.rstrip() + "\n" + addition, encoding="utf-8")


if __name__ == "__main__":
    main()
