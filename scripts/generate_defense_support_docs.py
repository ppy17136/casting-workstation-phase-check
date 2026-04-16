from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))


def main() -> None:
    artifacts_dir = WORKSPACE / "artifacts"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    speech_md = artifacts_dir / "砂型铸造本地智能工作站_5分钟答辩讲稿.md"
    speech_docx = artifacts_dir / "砂型铸造本地智能工作站_5分钟答辩讲稿.docx"
    qa_md = artifacts_dir / "砂型铸造本地智能工作站_答辩问答提纲.md"
    qa_docx = artifacts_dir / "砂型铸造本地智能工作站_答辩问答提纲.docx"
    manifest = artifacts_dir / "创新训练项目材料清单.md"

    speech_text = build_speech(summary)
    qa_text = build_qa(summary)

    speech_md.write_text(speech_text, encoding="utf-8")
    qa_md.write_text(qa_text, encoding="utf-8")
    build_docx("砂型铸造本地智能工作站 5 分钟答辩讲稿", speech_text, speech_docx)
    build_docx("砂型铸造本地智能工作站 答辩问答提纲", qa_text, qa_docx)
    update_manifest(manifest, speech_md, speech_docx, qa_md, qa_docx)


def build_speech(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站 5 分钟答辩讲稿",
            "",
            "各位老师好，我们的项目名称是“砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站”。",
            "本项目聚焦一个在铸造教学和竞赛中非常常见的问题：工艺图、工艺卡、三维模型、仿真结果和优化建议往往分散在不同软件和文件夹里，导致流程割裂、资料难复用、成果难展示。",
            "",
            "针对这个问题，我们设计了一套以本地桌面系统为核心的智能工作站。它不是纯网页系统，也不是工业 PLC 实时闭环控制系统，而是面向教学、竞赛和实验室管理场景的本地工程软件。",
            "系统的总体思路是，用统一的数据模型把项目管理、零件与材质、工艺方案、参数计算、SolidWorks 协同、ProCAST 仿真、工艺卡生成、AI 建议和成果导出连接起来，形成一个完整闭环。",
            "",
            "在技术路线方面，我们采用 PySide6 构建桌面界面，Python 作为本地业务核心，SQLite 作为本地数据库。",
            "SolidWorks 用于三维建模与 CAD 文件导出，ProCAST 用于仿真任务与结果管理，本系统则负责把这些专业软件中的数据组织起来，沉淀为项目资产。",
            "",
            "目前我们已经跑通了一套完整演示案例，项目编号是 "
            f"{summary['project_code']}，案例名称是“{summary['project_name']}”。",
            "在这个案例中，我们完成了以下闭环。",
            "第一，完成项目与零件基础信息录入，建立工艺方案版本。",
            "第二，基于零件和材料自动计算关键工艺参数，并保留公式来源。",
            "第三，关联本地 SolidWorks 三维文件，并记录导出结果。",
            "第四，登记 ProCAST 仿真任务与结果目录，实现结果统一管理。",
            "第五，自动生成工艺卡和质检/缺陷预防清单。",
            "第六，生成 AI 建议卡片，并保留证据链和人工审核状态。",
            "第七，最终自动打包输出成果包，便于课程展示、竞赛答辩和项目归档。",
            "",
            "这个系统的关键创新点主要有三方面。",
            "第一，是本地优先。它更适合实验室和教学现场，不依赖公网环境。",
            "第二，是一体化。它把原本分散的参数、模型、仿真、文档和建议统一到一个项目上下文中。",
            "第三，是可解释的 AI。我们没有把 AI 当成黑箱，而是采用“建议卡片 + 证据链 + 人工审核”的方式，使其更符合工程实践。",
            "",
            "从实际效果看，本次演示已经成功输出工艺卡、质检清单、"
            f"{summary['suggestion_count']} 条 AI 建议以及完整成果包。"
            f"其中当前仍有 {summary['pending_review_count']} 条建议保留在待审核状态，用于体现人工决策环节。",
            "",
            "下一步，我们计划进一步扩充材料库和规则库，接入真实大语言模型，增强 ProCAST 结果可视化能力，并逐步增加教师查看端和网页只读展示能力。",
            "",
            "总体来说，本项目已经具备可运行、可演示、可交付的阶段性成果，能够支撑创新训练项目的申报、中期检查和答辩展示。我的汇报完毕，谢谢各位老师。",
            "",
        ]
    )


def build_qa(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站答辩问答提纲",
            "",
            "## 1. 你们为什么做本地桌面系统，而不是纯网页？",
            "回答要点：",
            "- 教学实验室和竞赛环境中，很多电脑网络条件不稳定，本地系统更可靠。",
            "- SolidWorks、ProCAST 都是本地专业软件，本地桌面系统更容易协同集成。",
            "- 后续可以扩展网页只读展示，但核心工作流仍然保留在本地。",
            "",
            "## 2. 这个系统和一般工艺设计文档管理有什么区别？",
            "回答要点：",
            "- 不是简单的文件夹管理，而是把项目、零件、材质、工艺方案、参数、仿真、文档和建议统一建模。",
            "- 能自动生成工艺卡、质检清单和成果包，不只是存文件。",
            "- 能把 AI 建议接到工程流程里，并保留审核痕迹。",
            "",
            "## 3. AI 在系统中具体起什么作用？",
            "回答要点：",
            "- AI 不直接替代工程师决策，而是生成建议卡片。",
            "- 建议卡片包含建议内容、影响参数、证据来源和审核状态。",
            "- 当前演示案例已生成 "
            f"{summary['suggestion_count']} 条建议，其中保留 {summary['pending_review_count']} 条待审核建议。",
            "",
            "## 4. 为什么强调证据链和人工审核？",
            "回答要点：",
            "- 铸造工艺属于工程问题，不能完全依赖黑箱输出。",
            "- 建议必须能追溯来源，教师和学生都要知道“为什么这样改”。",
            "- 人工审核能保证系统适用于教学和工程试用场景。",
            "",
            "## 5. 你们为什么不做 PLC 实时闭环控制？",
            "回答要点：",
            "- 本项目定位是教学与设计辅助平台，不是产线自动控制系统。",
            "- 实时控制涉及硬件接口、安全和工业可靠性，超出当前创新训练项目边界。",
            "- 目前更聚焦设计、仿真、文档和决策支持闭环。",
            "",
            "## 6. 目前系统已经做到了什么程度？",
            "回答要点：",
            f"- 已跑通项目案例 {summary['project_code']}。",
            "- 页面链路已覆盖项目中心、零件与材质、工艺方案、参数计算、CAD 协同、仿真中心、结果管理、工艺卡、AI、审核和导出。",
            f"- 已生成工艺卡、清单、成果包 ZIP，成果包路径为：{summary['export_zip_path']}",
            "",
            "## 7. 这个系统后续还能扩展哪些方向？",
            "回答要点：",
            "- 增加更多铸造材料和经验规则库。",
            "- 接入真实本地或远程大语言模型。",
            "- 扩展教师查看端、网页只读展示和远程同步能力。",
            "- 增强 ProCAST 结果可视化和工艺图模板输出能力。",
            "",
            "## 8. 这个项目的创新点最核心是什么？",
            "回答要点：",
            "- 本地优先的一体化工作站形态。",
            "- 统一数据模型连接工艺、仿真、文档和 AI。",
            "- 证据化建议与人工审核机制。",
            "",
        ]
    )


def build_docx(title: str, text: str, output_path: Path) -> None:
    document = Document()
    lines = text.splitlines()
    if lines:
        document.add_heading(lines[0].lstrip("# ").strip(), level=0)
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(output_path)


def update_manifest(
    manifest_path: Path,
    speech_md: Path,
    speech_docx: Path,
    qa_md: Path,
    qa_docx: Path,
) -> None:
    content = manifest_path.read_text(encoding="utf-8")
    if "## 答辩辅助材料" in content:
        return
    addition = "\n".join(
        [
            "",
            "## 答辩辅助材料",
            f"- 5 分钟答辩讲稿 Markdown：{speech_md}",
            f"- 5 分钟答辩讲稿 Word：{speech_docx}",
            f"- 答辩问答提纲 Markdown：{qa_md}",
            f"- 答辩问答提纲 Word：{qa_docx}",
            "",
        ]
    )
    manifest_path.write_text(content.rstrip() + "\n" + addition, encoding="utf-8")


if __name__ == "__main__":
    main()
