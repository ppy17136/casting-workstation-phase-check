from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document
from docx.shared import Inches
from win32com.client import Dispatch


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))


PAGE_CAPTIONS = [
    ("dashboard", "仪表盘", "汇总项目数量、仿真任务、待审核建议与本地集成状态，作为教师查看与答辩演示的总入口。"),
    ("project_center", "项目中心", "完成项目编码、项目目录、零件基础信息与演示案例绑定，是整个流程的主索引页。"),
    ("parts", "零件与材质", "完整展示零件尺寸、壁厚、质量等级、热处理和材料热物性参数，支撑后续计算与工艺方案推导。"),
    ("scheme", "工艺方案", "展示分型方式、浇注位置、浇注系统类型和方案备注，实现工艺版本化管理。"),
    ("parameters", "参数计算", "集中呈现浇注时间、浇注重量、阻流截面积、收缩补偿等关键参数及其公式来源。"),
    ("solidworks", "SolidWorks 协同", "关联本地三维模型与导出文件，验证桌面系统与本地 CAD 的协同能力。"),
    ("simulation", "ProCAST 仿真中心", "登记仿真任务、输入输出目录和任务状态，支持课堂演示中的仿真闭环管理。"),
    ("results", "结果对比", "统一查看缩孔缩松、温度类结果索引和结果摘要，用于驱动工艺优化判断。"),
    ("documents", "工艺卡与质检", "自动生成工艺卡与质检/缺陷预防清单，把结构化数据转为比赛和答辩文档。"),
    ("ai", "AI 助手", "输出带证据链的建议卡片，体现 LLM/规则混合式辅助而非黑箱输出。"),
    ("review", "建议审核", "保留人工审批节点，支持通过/退回，满足工程审慎原则和教学考核需求。"),
    ("export", "成果导出", "自动收集工艺卡、CAD 文件、仿真结果和说明文件，生成成果包 ZIP。"),
    ("settings", "系统设置", "展示 LLM 配置、本地桥接路径、ProCAST 安装目录与默认项目目录，说明系统可部署性。"),
]


def main() -> None:
    artifacts_dir = WORKSPACE / "artifacts"
    walkthrough_dir = artifacts_dir / "walkthrough"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    project_name = summary.get("project_name") or "汽油机端盖一体化演示"
    project_code = summary.get("project_code") or "DEMO-K"

    markdown_path = artifacts_dir / "砂型铸造本地智能工作站_创新训练项目图文说明书.md"
    docx_path = artifacts_dir / "砂型铸造本地智能工作站_创新训练项目图文说明书.docx"
    outline_path = artifacts_dir / "砂型铸造本地智能工作站_创新训练项目答辩提纲.md"
    pptx_path = artifacts_dir / "砂型铸造本地智能工作站_创新训练项目答辩材料.pptx"
    manifest_path = artifacts_dir / "创新训练项目材料清单.md"

    markdown_path.write_text(build_markdown(summary), encoding="utf-8")
    outline_path.write_text(build_outline(summary), encoding="utf-8")
    build_docx(summary, walkthrough_dir, docx_path)
    build_ppt(summary, walkthrough_dir, pptx_path)
    manifest_path.write_text(
        "\n".join(
            [
                "# 创新训练项目材料清单",
                "",
                f"- 项目名称：{project_name}",
                f"- 项目编号：{project_code}",
                "",
                "## 正式材料",
                f"- 图文说明书 Markdown：{markdown_path}",
                f"- 图文说明书 Word：{docx_path}",
                f"- 答辩提纲：{outline_path}",
                f"- 答辩 PPT：{pptx_path}",
                f"- 操作视频：{artifacts_dir / 'demo_walkthrough.mp4'}",
                "",
                "## 关键截图",
                *[f"- {walkthrough_dir / f'{name}.png'}" for name, _, _ in PAGE_CAPTIONS],
                "",
            ]
        ),
        encoding="utf-8",
    )


def build_markdown(summary: dict) -> str:
    lines = [
        "# 砂型铸造本地智能工作站创新训练项目图文说明书",
        "",
        "## 一、项目背景",
        "传统砂型铸造教学与竞赛训练常见问题是：工艺图、工艺卡、三维模型、仿真结果和优化建议分散在不同软件与文件夹内，难以形成可追溯的项目闭环。",
        "本项目以本地桌面工作站为核心，围绕“工艺图—工艺卡—仿真辅助优化”建立统一的数据与流程平台，面向课程教学、创新训练、学科竞赛和实验室项目管理。",
        "",
        "## 二、建设目标",
        "1. 以本地桌面系统为主，支持在无公网或校园实验室环境中稳定运行。",
        "2. 优先兼容本地安装的 SolidWorks 和 ProCAST，实现 CAD/CAE 协同而不替代专业软件本身。",
        "3. 打通项目管理、参数计算、仿真结果管理、工艺卡生成、AI 建议与人工审核全链路。",
        "4. 预留远程同步、教师查看和网页只读展示接口，但 V1 不做实时闭环控制系统。",
        "",
        "## 三、系统总体设计",
        "系统采用 PySide6 桌面前端 + Python 本地业务核心 + SQLite 本地数据库 + SolidWorks/ProCAST 本地集成的组合架构。",
        "统一数据模型贯穿项目、零件、材质、工艺方案、参数、CAD 文件、仿真任务、文档、AI 建议和审批记录。",
        "",
        "### 架构要点",
        "- 本地优先：保证课堂、竞赛和实验室可离线运行。",
        "- 数据统一：避免“图纸一套、工艺卡一套、仿真目录一套”相互脱节。",
        "- 人工审核：AI 仅给建议卡片，不直接替代工程决策。",
        "- 可扩展：后续可增加教师查看端和远程同步能力。",
        "",
        "## 四、完整功能闭环",
        "系统已跑通一套完整演示任务，输出工艺卡、质检/缺陷预防清单、AI 建议卡片和成果包 ZIP。",
        "",
        f"- 项目编号：{summary['project_code']}",
        f"- 项目名称：{summary['project_name']}",
        f"- 项目目录：{summary['project_root']}",
        f"- 工艺卡：{summary['document_card_path']}",
        f"- 质检/缺陷预防清单：{summary['checklist_path']}",
        f"- AI 建议数：{summary['suggestion_count']}",
        f"- 待审核建议数：{summary['pending_review_count']}",
        f"- 成果包 ZIP：{summary['export_zip_path']}",
        "",
        "## 五、页面功能说明",
    ]

    for index, (_, title, description) in enumerate(PAGE_CAPTIONS, start=1):
        lines.extend([f"### {index}. {title}", description, ""])

    lines.extend(
        [
            "## 六、创新点总结",
            "1. 以本地桌面工作站作为核心载体，适合课程、竞赛和实验室联合使用。",
            "2. 将参数计算、CAD/CAE 协同、工艺卡输出和 AI 建议纳入同一项目上下文。",
            "3. AI 建议采用“建议卡片 + 证据链 + 人工审核”机制，强调工程可解释性。",
            "4. 支持把演示过程沉淀成成果包，便于项目申报、答辩和教学复用。",
            "",
            "## 七、后续实施计划",
            "1. 增加更多铸造材料与经验规则库，扩展参数计算覆盖面。",
            "2. 接入真实本地/远程 LLM 模型，替换当前演示级规则建议。",
            "3. 扩展 ProCAST 结果缩略图预览和结果对比分析页面。",
            "4. 增加教师查看端与网页只读展示能力，形成边缘—云协同模式。",
            "",
            "## 八、结论",
            "当前系统已经具备完整的本地可运行演示能力，可支撑创新训练项目申报、过程展示和阶段性答辩。后续重点应放在规则库扩充、LLM 接入和教学场景落地。",
            "",
        ]
    )
    return "\n".join(lines)


def build_outline(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站创新训练项目答辩提纲",
            "",
            "1. 项目背景与问题提出",
            "2. 研究目标与系统边界",
            "3. 总体架构与技术路线",
            "4. 演示案例与运行环境",
            "5. 各页面功能闭环展示",
            "6. AI 建议与人工审核机制",
            "7. 成果包导出与交付能力",
            "8. 创新点、应用价值与后续计划",
            "",
            f"- 演示项目：{summary['project_code']} / {summary['project_name']}",
            f"- 成果包：{summary['export_zip_path']}",
        ]
    )


def build_docx(summary: dict, walkthrough_dir: Path, output_path: Path) -> None:
    document = Document()
    document.add_heading("砂型铸造本地智能工作站创新训练项目图文说明书", level=0)
    document.add_paragraph("适用场景：创新训练项目申报、中期检查、答辩展示、课程演示。")

    for heading, paragraphs in (
        (
            "一、项目背景",
            [
                "传统砂型铸造工艺设计资料分散，项目过程难以复用和追踪。",
                "本项目围绕“工艺图—工艺卡—仿真辅助优化”建立本地优先的一体化工作站。",
            ],
        ),
        (
            "二、系统目标",
            [
                "以本地桌面工作站为主，保障实验室环境稳定运行。",
                "兼容 SolidWorks 与 ProCAST，形成 CAD/CAE 协同。",
                "实现项目管理、参数计算、文档输出、AI 建议和人工审核闭环。",
            ],
        ),
        (
            "三、演示案例摘要",
            [
                f"项目编号：{summary['project_code']}",
                f"项目名称：{summary['project_name']}",
                f"工艺卡：{summary['document_card_path']}",
                f"质检清单：{summary['checklist_path']}",
                f"AI 建议数：{summary['suggestion_count']}",
                f"成果包 ZIP：{summary['export_zip_path']}",
            ],
        ),
    ):
        document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph, style="List Bullet")

    document.add_heading("四、页面功能图解", level=1)
    for index, (name, title, description) in enumerate(PAGE_CAPTIONS, start=1):
        document.add_heading(f"{index}. {title}", level=2)
        document.add_paragraph(description)
        image_path = walkthrough_dir / f"{name}.png"
        if image_path.exists():
            document.add_picture(str(image_path), width=Inches(6.3))

    document.add_heading("五、创新点与应用价值", level=1)
    for item in (
        "本地优先，适合教学实验室和竞赛现场部署。",
        "统一管理工艺参数、三维模型、仿真结果和工艺文档。",
        "引入建议卡片、证据链和人工审核机制，提升可解释性。",
        "便于把运行过程直接沉淀为答辩和申报材料。",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("六、后续计划", level=1)
    for item in (
        "扩充材料库与工艺规则库。",
        "接入真实 LLM 和结果可视化分析。",
        "扩展教师查看端与网页只读展示能力。",
    ):
        document.add_paragraph(item, style="List Bullet")

    document.save(output_path)


def build_ppt(summary: dict, walkthrough_dir: Path, output_path: Path) -> None:
    app = Dispatch("PowerPoint.Application")
    app.Visible = 1
    presentation = app.Presentations.Add()

    def add_title_slide(title: str, subtitle: str) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 1)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = subtitle

    def add_bullet_slide(title: str, bullets: list[str]) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 2)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = "\r".join(bullets)

    def add_picture_slide(title: str, image_name: str, bullets: list[str]) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 2)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = "\r".join(bullets)
        image_path = walkthrough_dir / f"{image_name}.png"
        if image_path.exists():
            slide.Shapes.AddPicture(str(image_path), False, True, 420, 110, 500, 320)

    add_title_slide(
        "砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站",
        f"创新训练项目答辩材料\n案例：{summary['project_name']}",
    )
    add_bullet_slide(
        "项目背景与目标",
        [
            "面向砂型铸造教学、竞赛与实验室场景。",
            "解决项目资料分散、流程断裂、结果难复用的问题。",
            "形成“项目—参数—CAD/CAE—工艺卡—AI 建议—成果包”完整闭环。",
        ],
    )
    add_bullet_slide(
        "总体架构与技术路线",
        [
            "PySide6 桌面前端 + Python 本地核心 + SQLite 数据层。",
            "SolidWorks 负责 CAD，ProCAST 负责仿真，本系统负责组织、记录与输出。",
            "AI 模块采用建议卡片与人工审核机制，不做工业 PLC 实时控制。",
        ],
    )
    add_picture_slide(
        "项目与基础数据管理",
        "project_center",
        [
            "以项目中心作为主入口，统一管理项目编号、项目目录与零件信息。",
            "零件与材质页补充尺寸、壁厚、质量等级与材料热物性。",
        ],
    )
    add_picture_slide(
        "工艺方案与参数计算",
        "parameters",
        [
            "围绕当前方案自动计算关键参数并保留依据。",
            "支持后续与工艺卡、仿真任务和 AI 建议联动。",
        ],
    )
    add_picture_slide(
        "CAD/CAE 协同",
        "simulation",
        [
            "SolidWorks 页管理模型与导出结果。",
            "ProCAST 页管理仿真任务、输入输出目录和任务状态。",
        ],
    )
    add_picture_slide(
        "结果、文档与成果包",
        "documents",
        [
            "结果页集中查看仿真结果索引与摘要。",
            "文档页自动生成工艺卡、质检清单，导出页打包完整成果。",
        ],
    )
    add_picture_slide(
        "AI 建议与人工审核",
        "ai",
        [
            "AI 助手生成建议卡片并绑定证据来源。",
            "审核页保留人工通过/退回环节，保证工程可解释性。",
        ],
    )
    add_picture_slide(
        "系统设置与可部署性",
        "settings",
        [
            "配置 LLM 模式、桥接路径、ProCAST 安装目录和默认项目根目录。",
            "体现系统在实验室电脑上的真实部署方式。",
        ],
    )
    add_bullet_slide(
        "本次演示结果",
        [
            f"项目编号：{summary['project_code']}",
            f"工艺卡输出：{summary['document_card_path']}",
            f"AI 建议数：{summary['suggestion_count']}，待审核：{summary['pending_review_count']}",
            f"成果包 ZIP：{summary['export_zip_path']}",
        ],
    )
    add_bullet_slide(
        "创新点与后续计划",
        [
            "本地优先的一体化工作站形态，适合课堂、竞赛和实验室。",
            "统一数据模型连接工艺设计、仿真结果和 AI 建议。",
            "后续将接入真实 LLM、结果可视化与教师查看端。",
        ],
    )

    if output_path.exists():
        output_path.unlink()
    presentation.SaveAs(str(output_path))
    presentation.Close()
    app.Quit()


if __name__ == "__main__":
    main()
