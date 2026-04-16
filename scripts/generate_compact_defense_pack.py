from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document
from win32com.client import Dispatch


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))


def main() -> None:
    artifacts_dir = WORKSPACE / "artifacts"
    walkthrough_dir = artifacts_dir / "walkthrough"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    outline_md = artifacts_dir / "砂型铸造本地智能工作站_8页答辩版提纲.md"
    speaker_md = artifacts_dir / "砂型铸造本地智能工作站_逐页讲解词.md"
    speaker_docx = artifacts_dir / "砂型铸造本地智能工作站_逐页讲解词.docx"
    pptx_path = artifacts_dir / "砂型铸造本地智能工作站_8页答辩版.pptx"
    manifest_path = artifacts_dir / "创新训练项目材料清单.md"

    outline_text = build_outline(summary)
    speaker_text = build_speaker_notes(summary)

    outline_md.write_text(outline_text, encoding="utf-8")
    speaker_md.write_text(speaker_text, encoding="utf-8")
    build_docx("砂型铸造本地智能工作站逐页讲解词", speaker_text, speaker_docx)
    build_ppt(summary, walkthrough_dir, pptx_path)
    update_manifest(manifest_path, outline_md, speaker_md, speaker_docx, pptx_path)


def build_outline(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站 8 页答辩版提纲",
            "",
            "1. 课题背景与问题提出",
            "2. 研究目标与系统定位",
            "3. 总体架构与技术路线",
            "4. 项目与基础数据管理",
            "5. 参数计算与 CAD/CAE 协同",
            "6. 工艺卡、AI 建议与审核闭环",
            "7. 本次演示结果与成果包",
            "8. 创新点、价值与后续计划",
            "",
            f"- 演示项目：{summary['project_code']} / {summary['project_name']}",
            f"- 工艺卡：{summary['document_card_path']}",
            f"- 成果包 ZIP：{summary['export_zip_path']}",
        ]
    )


def build_speaker_notes(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站逐页讲解词",
            "",
            "## 第 1 页 课题背景与问题提出",
            "各位老师好，这一页主要说明我们为什么做这个系统。传统砂型铸造教学和竞赛中，工艺图、工艺卡、三维模型和仿真结果往往分散在不同软件里，资料难以统一管理，也不利于教学展示和项目复用。",
            "",
            "## 第 2 页 研究目标与系统定位",
            "我们的目标不是做工业 PLC 实时控制系统，而是做一个本地优先的智能工作站。它主要服务于课程教学、创新训练、竞赛和实验室项目管理，核心是把工艺设计、仿真和文档输出串成闭环。",
            "",
            "## 第 3 页 总体架构与技术路线",
            "系统采用 PySide6 桌面前端、Python 本地业务核心和 SQLite 本地数据库。SolidWorks 负责三维建模与导出，ProCAST 负责仿真，本系统负责组织、管理和输出工程数据。",
            "",
            "## 第 4 页 项目与基础数据管理",
            "这一页展示项目中心和零件与材质。项目中心统一管理项目编号、零件信息和目录；零件与材质页则补充尺寸、壁厚、质量等级和材料热物性，为后续计算和工艺方案提供数据基础。",
            "",
            "## 第 5 页 参数计算与 CAD/CAE 协同",
            "这一页强调的是工程计算和本地专业软件协同。系统可以自动计算关键工艺参数，保留公式依据，同时关联 SolidWorks 模型和 ProCAST 仿真任务，实现项目上下文统一。",
            "",
            "## 第 6 页 工艺卡、AI 建议与审核闭环",
            "这一页体现系统的智能化与工程审慎原则。系统能自动生成工艺卡和质检清单；AI 模块则输出建议卡片，但不会直接替代工程师决策，而是通过证据链和人工审核形成闭环。",
            "",
            "## 第 7 页 本次演示结果与成果包",
            f"本次演示项目编号是 {summary['project_code']}。目前已经成功输出工艺卡、质检清单、"
            f"{summary['suggestion_count']} 条 AI 建议，以及完整成果包 ZIP。"
            f"当前还保留 {summary['pending_review_count']} 条待审核建议，用来展示人工确认环节。",
            "",
            "## 第 8 页 创新点、价值与后续计划",
            "最后总结三点：第一，本地优先，适合实验室环境；第二，一体化，把参数、模型、仿真、文档和建议连接起来；第三，AI 可解释，通过建议卡片和审核机制保证工程可信。后续我们将接入真实大语言模型，扩展教师查看端和网页只读展示能力。",
            "",
        ]
    )


def build_docx(title: str, text: str, output_path: Path) -> None:
    document = Document()
    lines = text.splitlines()
    if lines:
        document.add_heading(lines[0].lstrip("# ").strip(), level=0)
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        document.add_paragraph(stripped)
    document.save(output_path)


def build_ppt(summary: dict, walkthrough_dir: Path, output_path: Path) -> None:
    app = Dispatch("PowerPoint.Application")
    app.Visible = 1
    presentation = app.Presentations.Add()

    def add_title_slide(title: str, subtitle: str) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 1)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = subtitle

    def add_bullet_slide(title: str, bullets: list[str], image_name: str | None = None) -> None:
        slide = presentation.Slides.Add(presentation.Slides.Count + 1, 2)
        slide.Shapes.Title.TextFrame.TextRange.Text = title
        slide.Shapes.Placeholders(2).TextFrame.TextRange.Text = "\r".join(bullets)
        if image_name:
            image_path = walkthrough_dir / f"{image_name}.png"
            if image_path.exists():
                slide.Shapes.AddPicture(str(image_path), False, True, 430, 110, 480, 300)

    add_title_slide(
        "砂型铸造工艺图—工艺卡—仿真辅助优化一体化本地智能工作站",
        f"8 页答辩版\n项目：{summary['project_name']}",
    )
    add_bullet_slide(
        "研究目标与系统定位",
        [
            "面向砂型铸造教学、竞赛和实验室管理。",
            "以本地桌面系统为主，不做 PLC 实时闭环控制。",
            "打通项目、参数、CAD/CAE、文档与 AI 建议闭环。",
        ],
    )
    add_bullet_slide(
        "总体架构与技术路线",
        [
            "PySide6 桌面前端 + Python 本地核心 + SQLite 数据库。",
            "SolidWorks 管 CAD，ProCAST 管仿真，本系统负责统一管理与输出。",
            "AI 模块采用建议卡片 + 证据链 + 人工审核机制。",
        ],
        "dashboard",
    )
    add_bullet_slide(
        "项目与基础数据管理",
        [
            "项目中心统一管理项目编号、目录、零件基础信息。",
            "零件与材质页完整维护尺寸、壁厚、质量等级和材料热物性。",
        ],
        "parts",
    )
    add_bullet_slide(
        "参数计算与 CAD/CAE 协同",
        [
            "自动计算关键工艺参数并保留计算依据。",
            "关联 SolidWorks 模型与 ProCAST 仿真任务，保证数据不脱节。",
        ],
        "simulation",
    )
    add_bullet_slide(
        "文档输出、AI 建议与审核",
        [
            "自动生成工艺卡和质检/缺陷预防清单。",
            "AI 生成建议卡片，审核页保留通过/退回和人工确认。",
        ],
        "ai",
    )
    add_bullet_slide(
        "本次演示结果",
        [
            f"项目编号：{summary['project_code']}",
            f"AI 建议数：{summary['suggestion_count']}，待审核：{summary['pending_review_count']}",
            f"成果包 ZIP：{summary['export_zip_path']}",
        ],
        "export",
    )
    add_bullet_slide(
        "创新点与后续计划",
        [
            "本地优先的一体化工作站形态。",
            "统一数据模型连接工艺、仿真、文档和 AI。",
            "后续接入真实 LLM、教师查看端和网页只读展示。",
        ],
        "settings",
    )

    if output_path.exists():
        output_path.unlink()
    presentation.SaveAs(str(output_path))
    presentation.Close()
    app.Quit()


def update_manifest(
    manifest_path: Path,
    outline_md: Path,
    speaker_md: Path,
    speaker_docx: Path,
    pptx_path: Path,
) -> None:
    content = manifest_path.read_text(encoding="utf-8")
    if "## 精简答辩版材料" in content:
        return
    addition = "\n".join(
        [
            "",
            "## 精简答辩版材料",
            f"- 8 页答辩版提纲：{outline_md}",
            f"- 8 页答辩版 PPT：{pptx_path}",
            f"- 逐页讲解词 Markdown：{speaker_md}",
            f"- 逐页讲解词 Word：{speaker_docx}",
            "",
        ]
    )
    manifest_path.write_text(content.rstrip() + "\n" + addition, encoding="utf-8")


if __name__ == "__main__":
    main()
