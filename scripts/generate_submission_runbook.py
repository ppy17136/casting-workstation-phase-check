from __future__ import annotations

import json
from pathlib import Path
import sys

from docx import Document


WORKSPACE = Path(__file__).resolve().parents[1]
if str(WORKSPACE) not in sys.path:
    sys.path.insert(0, str(WORKSPACE))


def main() -> None:
    artifacts_dir = WORKSPACE / "artifacts"
    summary = json.loads((artifacts_dir / "demo_workflow_summary.json").read_text(encoding="utf-8"))

    checklist_md = artifacts_dir / "砂型铸造本地智能工作站_最终提交清单.md"
    checklist_docx = artifacts_dir / "砂型铸造本地智能工作站_最终提交清单.docx"
    runbook_md = artifacts_dir / "砂型铸造本地智能工作站_答辩当天操作手册.md"
    runbook_docx = artifacts_dir / "砂型铸造本地智能工作站_答辩当天操作手册.docx"
    manifest_path = artifacts_dir / "创新训练项目材料清单.md"

    checklist_text = build_checklist(summary)
    runbook_text = build_runbook(summary)

    checklist_md.write_text(checklist_text, encoding="utf-8")
    runbook_md.write_text(runbook_text, encoding="utf-8")
    build_docx("砂型铸造本地智能工作站最终提交清单", checklist_text, checklist_docx)
    build_docx("砂型铸造本地智能工作站答辩当天操作手册", runbook_text, runbook_docx)
    update_manifest(manifest_path, checklist_md, checklist_docx, runbook_md, runbook_docx)


def build_checklist(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站最终提交清单",
            "",
            "## 一、软件与安装交付",
            "- 安装包：CastingWorkstationSetup.exe",
            "- 便携版：CastingWorkstation.exe 所在目录",
            "- SolidWorksBridge 发布目录",
            "- 安装说明或 README",
            "",
            "## 二、项目演示材料",
            f"- 演示项目编号：{summary['project_code']}",
            f"- 项目名称：{summary['project_name']}",
            f"- 工艺卡：{summary['document_card_path']}",
            f"- 质检/缺陷预防清单：{summary['checklist_path']}",
            f"- 成果包 ZIP：{summary['export_zip_path']}",
            "",
            "## 三、答辩与申报材料",
            "- 创新训练项目图文说明书（正式版）",
            "- 创新训练项目答辩材料 PPT（完整版）",
            "- 8 页答辩版 PPT（精简版）",
            "- 5 分钟答辩讲稿",
            "- 3 分钟极简汇报稿",
            "- 答辩问答提纲",
            "- 逐页讲解词",
            "- 申请书正文素材",
            "",
            "## 四、演示支撑文件",
            "- demo_walkthrough.mp4 操作视频",
            "- walkthrough 截图目录",
            "- 创新训练项目材料清单",
            "",
            "## 五、建议打包方式",
            "- 文件夹 1：软件安装包与便携版",
            "- 文件夹 2：演示案例与成果包",
            "- 文件夹 3：说明书、PPT、讲稿、问答提纲",
            "- 文件夹 4：视频与关键截图",
            "",
            "## 六、答辩前最后核对",
            "- 软件是否可正常启动",
            "- SolidWorks 与 ProCAST 是否可识别",
            "- 项目演示数据是否已加载",
            "- 工艺卡、成果包路径是否能打开",
            "- PPT 是否为最终版本",
            "- 视频和截图是否可正常播放或查看",
            "",
        ]
    )


def build_runbook(summary: dict) -> str:
    return "\n".join(
        [
            "# 砂型铸造本地智能工作站答辩当天操作手册",
            "",
            "## 一、答辩前 30 分钟",
            "- 打开电脑并确认电源、投影和鼠标正常。",
            "- 打开便携版或安装版软件，确认主界面能正常进入。",
            "- 检查 SolidWorks 是否可启动，ProCAST 路径是否识别。",
            "- 提前打开答辩 PPT 和 5 分钟讲稿。",
            "",
            "## 二、建议提前打开的文件",
            "- 完整版答辩 PPT",
            "- 8 页答辩版 PPT",
            "- 5 分钟答辩讲稿",
            "- 答辩问答提纲",
            "- 申请书正文素材",
            f"- 工艺卡文件：{summary['document_card_path']}",
            f"- 成果包 ZIP：{summary['export_zip_path']}",
            "",
            "## 三、现场演示推荐顺序",
            "1. 先讲 PPT 前 3 页：背景、目标、架构。",
            "2. 打开软件演示项目中心与零件材质页。",
            "3. 切到参数计算页，说明系统如何自动计算关键参数。",
            "4. 切到 SolidWorks / ProCAST 页面，说明与本地专业软件协同。",
            "5. 切到工艺卡页面，展示文档自动生成能力。",
            "6. 切到 AI 助手和建议审核页面，强调证据链与人工审核。",
            "7. 切到成果导出页面，说明最终可形成完整成果包。",
            "8. 回到 PPT 最后 2 页，总结创新点和后续计划。",
            "",
            "## 四、时间分配建议",
            "- 3 分钟场景：只讲背景、架构、闭环和创新点。",
            "- 5 分钟场景：讲完整套闭环，并简短演示 3 到 4 个关键页面。",
            "- 8 分钟以上场景：加入工艺卡、AI 建议和成果包展示。",
            "",
            "## 五、老师可能现场要求你点开的内容",
            "- 项目中心：证明系统不是静态截图，而是真实项目数据。",
            "- 零件与材质：证明页面内容完整，不是空壳。",
            "- 参数计算：证明有工程计算逻辑。",
            "- 工艺卡：证明能输出实际文档。",
            "- AI 建议与审核：证明有智能辅助但保留人工决策。",
            "- 成果导出：证明有交付能力。",
            "",
            "## 六、现场风险与应对",
            "- 如果软件启动慢：先讲 PPT，不要空等。",
            "- 如果 SolidWorks 未响应：直接展示已关联结果和导出文件。",
            "- 如果 ProCAST 无法临时运行：展示仿真任务登记和结果归档页即可。",
            "- 如果时间被压缩：立即切换到 3 分钟极简汇报稿。",
            "",
            "## 七、结束语模板",
            "本项目已经形成可运行、可演示、可交付的阶段性成果，能够支撑创新训练项目申报、教学展示和竞赛答辩。下一步将继续增强规则库、LLM 接入和教师查看能力。谢谢各位老师。",
            "",
        ]
    )


def build_docx(title: str, text: str, output_path: Path) -> None:
    document = Document()
    lines = text.splitlines()
    if lines:
        document.add_heading(lines[0].lstrip("# ").strip(), level=0)
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(output_path)


def update_manifest(
    manifest_path: Path,
    checklist_md: Path,
    checklist_docx: Path,
    runbook_md: Path,
    runbook_docx: Path,
) -> None:
    content = manifest_path.read_text(encoding="utf-8")
    if "## 提交与现场演示材料" in content:
        return
    addition = "\n".join(
        [
            "",
            "## 提交与现场演示材料",
            f"- 最终提交清单 Markdown：{checklist_md}",
            f"- 最终提交清单 Word：{checklist_docx}",
            f"- 答辩当天操作手册 Markdown：{runbook_md}",
            f"- 答辩当天操作手册 Word：{runbook_docx}",
            "",
        ]
    )
    manifest_path.write_text(content.rstrip() + "\n" + addition, encoding="utf-8")


if __name__ == "__main__":
    main()
