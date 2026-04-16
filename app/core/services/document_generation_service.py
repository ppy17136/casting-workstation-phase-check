from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document

from app.core.repositories.document_repository import DocumentRepository
from app.core.repositories.parameter_repository import ParameterRecord
from app.core.repositories.part_repository import PartRecord
from app.core.repositories.project_repository import ProjectRecord
from app.core.repositories.scheme_repository import SchemeRecord


@dataclass(slots=True)
class GeneratedDocumentBundle:
    card_path: Path
    checklist_path: Path
    markdown_card_path: Path
    card_no: str


class DocumentGenerationService:
    def __init__(self, repository: DocumentRepository) -> None:
        self.repository = repository

    def generate_bundle(
        self,
        *,
        project: ProjectRecord,
        part: PartRecord,
        scheme: SchemeRecord,
        parameters: list[ParameterRecord],
    ) -> GeneratedDocumentBundle:
        docs_dir = Path(project.root_dir) / "generated_docs"
        docs_dir.mkdir(parents=True, exist_ok=True)

        card_no = f"{project.project_code}-{scheme.scheme_code}-{scheme.version_no}"
        card_path = docs_dir / f"{card_no}_process_card.docx"
        markdown_card_path = docs_dir / f"{card_no}_process_card.md"
        checklist_path = docs_dir / f"{card_no}_inspection_checklist.md"

        self._write_markdown_card(markdown_card_path, project, part, scheme, parameters, card_no)
        self._write_docx_card(card_path, project, part, scheme, parameters, card_no)

        checklist_items = self._build_checklist_items(part, scheme, parameters)
        checklist_content = "\n".join(
            [
                f"# 缺陷预防与质检清单 {card_no}",
                "",
                "| 类别 | 项目 | 控制阶段 | 控制方法 | 验收规则 | 风险说明 |",
                "| --- | --- | --- | --- | --- | --- |",
                *[
                    f"| {item['item_type']} | {item['item_name']} | {item['control_stage']} | {item['control_method']} | {item['acceptance_rule']} | {item['risk_reason']} |"
                    for item in checklist_items
                ],
                "",
            ]
        )
        checklist_path.write_text(checklist_content, encoding="utf-8")

        self.repository.create_process_card(
            scheme_id=scheme.id,
            card_no=card_no,
            template_name="docx_v1",
            docx_path=str(card_path),
            pdf_path="",
        )
        self.repository.replace_inspection_items(scheme.id, checklist_items)

        return GeneratedDocumentBundle(
            card_path=card_path,
            checklist_path=checklist_path,
            markdown_card_path=markdown_card_path,
            card_no=card_no,
        )

    def _write_markdown_card(
        self,
        card_path: Path,
        project: ProjectRecord,
        part: PartRecord,
        scheme: SchemeRecord,
        parameters: list[ParameterRecord],
        card_no: str,
    ) -> None:
        parameter_lines = [
            f"| {item.param_group} | {item.param_name} | {item.param_value} | {item.param_unit} | {item.source_type or item.value_type} |"
            for item in parameters
        ]
        if not parameter_lines:
            parameter_lines.append("| - | 暂无参数 | - | - | - |")

        card_content = "\n".join(
            [
                f"# 工艺卡 {card_no}",
                "",
                "## 项目概况",
                f"- 项目编号：{project.project_code}",
                f"- 项目名称：{project.project_name}",
                f"- 铸造方式：{project.casting_method or '-'}",
                f"- 零件名称：{part.part_name}",
                f"- 图号：{part.drawing_no or '-'}",
                f"- 材料：{part.material_name or '-'}",
                "",
                "## 工艺方案",
                f"- 方案编号：{scheme.scheme_code}",
                f"- 方案名称：{scheme.scheme_name}",
                f"- 版本：{scheme.version_no}",
                f"- 型腔形式：{scheme.mold_type or '-'}",
                f"- 分型方式：{scheme.parting_method or '-'}",
                f"- 浇注位置：{scheme.pouring_position or '-'}",
                f"- 浇注系统：{scheme.gating_type or '-'}",
                "",
                "## 关键工艺参数",
                "| 分组 | 参数 | 数值 | 单位 | 来源 |",
                "| --- | --- | --- | --- | --- |",
                *parameter_lines,
                "",
                "## 备注",
                scheme.notes or "-",
                "",
            ]
        )
        card_path.write_text(card_content, encoding="utf-8")

    def _write_docx_card(
        self,
        card_path: Path,
        project: ProjectRecord,
        part: PartRecord,
        scheme: SchemeRecord,
        parameters: list[ParameterRecord],
        card_no: str,
    ) -> None:
        document = Document()
        document.add_heading(f"工艺卡 {card_no}", level=0)

        document.add_heading("项目概况", level=1)
        for label, value in (
            ("项目编号", project.project_code),
            ("项目名称", project.project_name),
            ("铸造方式", project.casting_method or "-"),
            ("零件名称", part.part_name),
            ("图号", part.drawing_no or "-"),
            ("材料", part.material_name or "-"),
        ):
            document.add_paragraph(f"{label}：{value}")

        document.add_heading("工艺方案", level=1)
        for label, value in (
            ("方案编号", scheme.scheme_code),
            ("方案名称", scheme.scheme_name),
            ("版本", scheme.version_no),
            ("型腔形式", scheme.mold_type or "-"),
            ("分型方式", scheme.parting_method or "-"),
            ("浇注位置", scheme.pouring_position or "-"),
            ("浇注系统", scheme.gating_type or "-"),
        ):
            document.add_paragraph(f"{label}：{value}")

        document.add_heading("关键工艺参数", level=1)
        table = document.add_table(rows=1, cols=5)
        header = table.rows[0].cells
        for index, title in enumerate(("分组", "参数", "数值", "单位", "来源")):
            header[index].text = title
        for item in parameters:
            row = table.add_row().cells
            row[0].text = item.param_group
            row[1].text = item.param_name
            row[2].text = item.param_value
            row[3].text = item.param_unit
            row[4].text = item.source_type or item.value_type
        if not parameters:
            row = table.add_row().cells
            row[0].text = "-"
            row[1].text = "暂无参数"
            row[2].text = "-"
            row[3].text = "-"
            row[4].text = "-"

        document.add_heading("备注", level=1)
        document.add_paragraph(scheme.notes or "-")
        document.save(card_path)

    def _build_checklist_items(
        self,
        part: PartRecord,
        scheme: SchemeRecord,
        parameters: list[ParameterRecord],
    ) -> list[dict[str, str]]:
        has_feeding = any(item.param_group == "feeding" for item in parameters)
        return [
            {
                "item_type": "process",
                "item_name": "分型与起模检查",
                "control_stage": "造型前",
                "control_method": "核对分型方式与起模斜度",
                "acceptance_rule": scheme.parting_method or "与工艺方案一致",
                "risk_reason": "分型不当会造成错箱、起模损伤和尺寸偏差",
            },
            {
                "item_type": "process",
                "item_name": "浇注系统复核",
                "control_stage": "制芯/合箱前",
                "control_method": "核对浇口位置、内浇道截面与阻流面积",
                "acceptance_rule": scheme.gating_type or "与方案一致",
                "risk_reason": "浇注系统不匹配会导致紊流、卷气和充型不足",
            },
            {
                "item_type": "quality",
                "item_name": "壁厚热节检查",
                "control_stage": "工艺评审",
                "control_method": "结合最大壁厚和冒口设置复核补缩路径",
                "acceptance_rule": "热节区域需具备有效补缩",
                "risk_reason": "厚大截面和热节区域易出现缩孔缩松",
            },
            {
                "item_type": "quality",
                "item_name": "重量与出品率校验",
                "control_stage": "方案定稿",
                "control_method": "复核净重、毛坯重与出品率",
                "acceptance_rule": f"净重 {part.net_weight or 0:.3f} kg，毛坯重 {part.blank_weight or 0:.3f} kg",
                "risk_reason": "重量数据异常会直接影响浇注时间、阻流面积与成本估算",
            },
            {
                "item_type": "quality",
                "item_name": "补缩系统确认",
                "control_stage": "方案定稿",
                "control_method": "检查冒口、冷铁或保温措施",
                "acceptance_rule": "存在补缩措施" if has_feeding else "需补录补缩措施",
                "risk_reason": "缺少补缩措施时高风险区域更易产生缩松缺陷",
            },
        ]
